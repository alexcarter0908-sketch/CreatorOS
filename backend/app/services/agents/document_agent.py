from __future__ import annotations

import base64
from io import BytesIO

from docx import Document as DocxDocument

from app.schemas.ai_request import AIRequest
from app.schemas.ai_response import AIResponse
from app.services.agents.base_agent import BaseAgent


_DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def _extract_title(raw_text: str) -> str:
    for line in raw_text.split("\n"):
        line = line.strip()
        if line.startswith("# "):
            return line[2:].strip()
    return "Document"


def _build_docx_bytes(raw_text: str) -> bytes:
    """
    Converts the LLM's generated content into an actual .docx file. The
    model is instructed (see the prompt in execute()) to mark the
    document title with a leading '# ' and section headings with '## ',
    so this parser turns them into real Word heading styles instead of
    leaving literal '#' characters in the document.
    """
    doc = DocxDocument()

    for raw_line in raw_text.split("\n"):
        line = raw_line.strip()
        if not line:
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        else:
            doc.add_paragraph(line)

    stream = BytesIO()
    doc.save(stream)
    return stream.getvalue()


class DocumentAgent(BaseAgent):
    """
    Generates real, downloadable Word (.docx) documents - client
    pitches, proposals, reports, explainers - instead of only replying
    with plain chat text. Content is written by the same reliable
    text/Groq pipeline every other agent uses, then converted into an
    actual Word file and stored like any other generated asset (image/
    video/audio), so it shows up with a real download link.
    """

    name = "document"

    async def execute(
        self,
        request: AIRequest,
    ) -> AIResponse:
        original_prompt = request.prompt

        request.asset_type = "text"
        request.prompt = f"""
You are an expert business writer producing a polished, client-ready
document (a pitch, proposal, report, or explainer - whichever best fits
the request below).

Write complete, professional, well-organized content. Structure it with
a single document title and clear section headings.

FORMAT RULES (strict - this output is parsed into a real Word document,
not shown as chat text):
- The very first line must be the document title, prefixed with "# "
  (a single '#' followed by a space), e.g. "# CreatorOS Client Pitch"
- Each section heading must be on its own line, prefixed with "## "
  (two '#' characters followed by a space), e.g. "## What Is CreatorOS"
- Every other line is normal paragraph text - do NOT use any other
  markdown symbols anywhere (no **, no *, no numbered-list dashes) -
  write plain sentences; if you need a short list, write each item as
  its own plain line instead of a markdown bullet.
- Do not use the [Xs-Ys] timed-beat format used for video scripts -
  this is a written document, not a video script.

User Request:
{original_prompt}
"""
        request.metadata.update(
            {
                "pipeline": "document_generation",
                "agent": self.name,
            }
        )

        generated = await self.generate(request)
        raw_text = generated.get("result", "") if isinstance(generated, dict) else str(generated)
        if isinstance(raw_text, dict):
            raw_text = raw_text.get("text", "")
        raw_text = str(raw_text)

        docx_bytes = _build_docx_bytes(raw_text)
        encoded = base64.b64encode(docx_bytes).decode("ascii")
        data_uri = f"data:{_DOCX_MIME};base64,{encoded}"

        provider_metadata = generated.get("metadata", {}) if isinstance(generated, dict) else {}
        title = _extract_title(raw_text)

        return {
            "result": {"url": data_uri},
            "metadata": {
                **provider_metadata,
                "text": raw_text.strip() or f"Document ready: {title}",
            },
            "provider": "creatoros",
            "model": "docx-generator",
        }

